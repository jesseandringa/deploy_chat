import os
import re

# import fitz
import fitz
import mistral
import openpyxl
from docx import Document


class DocumentTracker:
    def __init__(self):
        self.chapter = None
        self.article = None
        self.section = None
        self.division = None
        self.chapter_name = None
        self.article_name = None
        self.section_title = None
        self.division_name = None

    def update_chapter(self, chapter, chapter_name=None):
        self.chapter = chapter
        self.chapter_name = chapter_name
        self.article = None  # Reset article, section, and division when chapter changes
        self.section = None
        self.division = None

    def update_article(self, article, article_name=None):
        self.article = article
        self.article_name = article_name
        self.section = None  # Reset section and division when article changes
        self.division = None

    def update_division(self, division, division_name=None):
        self.division = division
        self.division_name = division_name
        self.section = None  # Reset section when division changes

    def update_section(self, section, section_title=None):
        self.section = section
        self.section_title = section_title

    def _abbreviate(self, phrase, max_length=4):
        """
        Abbreviate the given phrase by taking the first two letters of each word,
        skipping common words like 'and', 'or', 'the', until the max_length is reached.
        """
        if not phrase:
            return ""
        common_words = {
            "and",
            "or",
            "the",
            "of",
            "in",
            "on",
            "to",
            "for",
            "with",
            "at",
            "by",
        }
        words = re.split(r"\W+", phrase)  # Split on non-alphanumeric characters
        abbreviation = "".join(
            word[:2].upper()
            for word in words
            if word and word.lower() not in common_words
        )
        return abbreviation  # [:max_length]  # Limit to max_length characters

    def _format_section(self, section):
        """
        Formats section numbers like '2-1' to 'S2-1'
        """
        if not section:
            return ""
        return f"S{section.replace('Sec', '').replace(' ', '').replace('.', '')}"

    def parse_input(self, input_string):
        """
        Parses a string like 'Chapter 2 ADMINISTRATION', 'DIVISION 1. GENERALLY' and updates the chapter, article, or division.
        """
        # Match 'Chapter X NAME'
        chapter_match = re.match(r"Chapter\s+(\d+)\s+(.+)", input_string, re.IGNORECASE)
        if chapter_match:
            chapter_num = chapter_match.group(1)
            chapter_name = chapter_match.group(2)
            self.update_chapter(chapter_num, chapter_name)
            return

        # Match 'Article X NAME'
        article_match = re.match(
            r"Article\s+(\w+)\.\s+(.+)", input_string, re.IGNORECASE
        )
        if article_match:
            article_num = article_match.group(1)
            article_name = article_match.group(2)
            self.update_article(article_num, article_name)
            return

        # Match 'Sec X-X TITLE'
        section_match = re.match(
            r"Sec\.\s+([\d\-]+)\.\s+(.+)", input_string, re.IGNORECASE
        )
        if section_match:
            section_num = section_match.group(1)
            section_title = section_match.group(2)
            self.update_section(section_num, section_title)
            return

        # Match 'Division X NAME'
        division_match = re.match(
            r"DIVISION\s+(\d+)\.\s+(.+)", input_string, re.IGNORECASE
        )
        if division_match:
            division_num = division_match.group(1)
            division_name = division_match.group(2)
            self.update_division(division_num, division_name)

    def generate_url_route(self):
        """
        Generate a custom URL route from the current chapter, division, article, and section details.
        For example: CH2AD_ARTI_DIV1GE_S2-1INFOGOCL (only include division if it exists).
        """
        if not self.chapter:
            return ""

        chapter_abbr = self._abbreviate(self.chapter_name) if self.chapter_name else ""
        article_abbr = self._abbreviate(self.article_name) if self.article_name else ""
        section_abbr = (
            self._abbreviate(self.section_title, max_length=8)
            if self.section_title
            else ""
        )
        section_code = self._format_section(self.section)
        division_abbr = (
            self._abbreviate(self.division_name) if self.division_name else ""
        )
        division_code = f"_DIV{self.division}{division_abbr}" if self.division else ""
        article_code = f"_ART{self.article}{article_abbr}" if self.article else ""
        # Construct the encoded URL route
        return f"CH{self.chapter}{chapter_abbr}{article_code}{division_code}_{section_code}{section_abbr}"

    def current_location(self):
        return f"Chapter {self.chapter}, Division {self.division}, Article {self.article}, Section {self.section}"


def chunk_municode_doc_into_paragraphs(doc_path, chunk_size=500):
    """Chunks a Municode DOC file into paragraphs."""
    doc_tracker = DocumentTracker()
    chunks = []

    # Open the DOCX file
    try:
        doc = Document(doc_path)
    except Exception:
        return []

    paragraphs_and_sources = []
    for para in doc.paragraphs:
        if para.text.lower().startswith("chapter "):
            doc_tracker.parse_input(para.text)
        if para.text.lower().startswith("article "):
            doc_tracker.parse_input(para.text)
        if para.text.lower().startswith("division "):
            doc_tracker.parse_input(para.text)
        if para.text.lower().startswith("sec. "):
            doc_tracker.parse_input(para.text)
        url = doc_tracker.generate_url_route()
        paragraphs_and_sources.append([para.text, url])

        # print(doc.paragraphs[i].text)
        # print("::::\n")

    # url = doc_tracker.generate_url_route()
    # print(url)
    # print(doc_tracker.current_location())

    # text_chunks = [text[i : i + chunk_size] for i in range(0, len(text), chunk_size)]
    text = ""
    for ps in paragraphs_and_sources:
        sources = set()

        if text == "":
            text = ps[0]
        else:
            text = text + " " + ps[0]
        sources.add(ps[1])
        if len(text) > chunk_size:
            text = text.replace("\ax0", " ")
            chunks.append([sources, text])
            text = ""

    return chunks


def chunk_file_into_paragraphs(file_path, chunk_size=500, use_llm=False):
    """Chunks a file into paragraphs."""
    # print out converted filename
    converted_file_path = file_path.replace("$$$", "/")
    print(f"converted file path: {converted_file_path}")
    if file_path.endswith(".pdf"):
        return chunk_pdf_into_paragraphs(file_path, chunk_size, use_llm)
    elif (
        file_path.endswith(".xlsx")
        or file_path.endswith(".xls")
        or file_path.endswith(".xlsm")
        or file_path.endswith(".csv")
    ):
        return chunk_xlsx_into_paragraphs(file_path, use_llm)
    elif file_path.endswith(".docx") or file_path.endswith(".doc"):
        return chunk_docx_into_paragraphs(file_path, chunk_size, use_llm)
    else:
        return []


def get_text_from_file(file_path):
    """Get text from a file."""
    if file_path.endswith(".pdf"):
        return get_text_from_pdf(file_path)
    elif (
        file_path.endswith(".xlsx")
        or file_path.endswith(".xls")
        or file_path.endswith(".xlsm")
        or file_path.endswith(".csv")
    ):
        return get_text_from_xlsx(file_path)
    elif file_path.endswith(".docx") or file_path.endswith(".doc"):
        return get_text_from_docx(file_path)
    else:
        return [file_path, ""]


def get_text_from_pdf(pdf_path):
    """Get text from a PDF."""
    try:
        # Open the PDF file
        pdf_document = fitz.open(pdf_path)

        # Extract text from each page
        text = ""
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            text += page.get_text("text")

        return [pdf_path, text]
    except Exception:
        return [pdf_path, ""]


def get_text_from_xlsx(xlsx_path):
    """Get text from an XLSX file."""
    try:
        text = ""

        # Open the XLSX file
        xlsx_document = openpyxl.load_workbook(xlsx_path)
        sheet = xlsx_document.active

        for row in sheet.iter_rows(
            min_row=1, max_row=sheet.max_row, max_col=sheet.max_column
        ):
            for cell in row:
                if cell.value:
                    text += str(cell.value) + " "

        return [xlsx_path, text]
    except Exception:
        return [xlsx_path, ""]


def get_text_from_docx(docx_path):
    """Get text from a DOCX file."""
    try:
        doc = Document(docx_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return [docx_path, text]
    except Exception:
        return [docx_path, "text"]


def chunk_xlsx_into_paragraphs(xlsx_path):
    """Chunks an XLSX file into paragraphs."""
    try:
        chunks = []

        # Open the XLSX file
        xlsx_document = openpyxl.load_workbook(xlsx_path)
        sheet = xlsx_document.active

        paragraph_number = 0
        labels = []
        for row in sheet.iter_rows(
            min_row=1, max_row=sheet.max_row, max_col=sheet.max_column
        ):
            paragraph = ""
            for i, cell in enumerate(row):
                if paragraph_number == 0:
                    labels.append(cell.value)
                elif cell.value:
                    paragraph += str(labels[i]) + ": " + str(cell.value) + ", "
            if paragraph_number > 0:
                chunk = [xlsx_path, paragraph]
                chunks.append(chunk)
            paragraph_number += 1
        return chunks
    except Exception:
        return []


def chunk_pdf_into_paragraphs(pdf_path, chunk_size=500, use_llm=False):
    """Chunks a PDF into paragraphs."""
    chunks = []

    # Open the PDF file
    try:
        pdf_document = fitz.open(pdf_path)
    except Exception:
        return []

    paragraph_number = 0
    all_text = ""
    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        text = page.get_text("text")
        all_text += text
    text_chunks = [
        all_text[i : i + chunk_size] for i in range(0, len(all_text), chunk_size)
    ]
    for p in text_chunks:
        if use_llm:
            is_usable = mistral.check_if_text_is_helpful_using_ollama(p)
            if is_usable:
                cleaned_chunk = mistral.trim_text_using_ollama(p)
            print(f"\n\n\nis usable text: {is_usable}")
            print(cleaned_chunk)
        else:
            cleaned_chunk = p.replace("\ax0", "")

        chunk = [pdf_path, cleaned_chunk]
        chunks.append(chunk)
    return chunks


def get_all_pdfs(folder_path):
    """Get all PDF files in a folder."""
    # List all files and directories in the given folder
    items = os.listdir(folder_path)

    # Filter out the PDFs
    pdfs = [
        item
        for item in items
        if item.endswith(".pdf") and os.path.isfile(os.path.join(folder_path, item))
    ]

    return pdfs


def get_all_files_with_full_path(folder_path):
    """Get all files in a folder."""
    # List all files and directories in the given folder
    items = os.listdir(folder_path)

    # Filter out the PDFs
    files = sorted(
        os.path.join(folder_path, item)
        for item in items
        if os.path.isfile(os.path.join(folder_path, item))
    )

    return files


def get_all_xlsx(folder_path):
    """Get all XLSX files in a folder."""
    xlsx_files = []
    for root, _, files in os.walk(folder_path):
        for file in files:
            if file.endswith(".xlsx"):
                xlsx_files.append(os.path.join(root, file))
    return xlsx_files


def get_all_docs(folder_path):
    """Get all DOCX files in a folder."""
    doc_files = []
    for root, _, files in os.walk(folder_path):
        for file in files:
            if file.endswith(".docx"):
                doc_files.append(os.path.join(root, file))
    return doc_files


def chunk_docx_into_paragraphs(docx_path, chunk_size=500, use_llm=False):
    """Chunks a DOCX file into paragraphs."""
    chunks = []

    # Open the DOCX file
    try:
        doc = Document(docx_path)
    except Exception:
        return []
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    text = remove_initial_all_caps(text)
    text_chunks = [text[i : i + chunk_size] for i in range(0, len(text), chunk_size)]
    for p in text_chunks:
        if use_llm:
            is_usable = mistral.check_if_text_is_helpful_using_ollama(p)
            if is_usable:
                cleaned_chunk = mistral.trim_text_using_ollama(p)
                chunk = [docx_path, cleaned_chunk]
                chunks.append(chunk)
                print("usable")
                print(cleaned_chunk)
            else:
                print("not usable")
                print(p)

    return chunks


def check_validity_of_file(file_path, chunks=None):
    if chunks is None:
        if file_path.endswith(".pdf"):
            chunks = chunk_pdf_into_paragraphs(file_path)
        elif file_path.endswith(".xlsx"):
            chunks = chunk_xlsx_into_paragraphs(file_path)
        elif file_path.endswith(".docx"):
            chunks = chunk_docx_into_paragraphs(file_path)

    total_text = ""

    for chunk in chunks:
        total_text += chunk[1]

    if is_mostly_binary(total_text):
        return 0
    if text_is_too_short(total_text):
        return 0
    return len(total_text)


def text_is_too_short(text, threshold=100):
    return len(text) < threshold


def is_mostly_binary(text, threshold=0.5):
    # Count non-printable characters
    non_printable_count = sum(
        1 for char in text if not (32 <= ord(char) <= 126 or char in "\n\r\t")
    )

    # Calculate the proportion of non-printable characters
    if len(text) == 0:
        return True
    binary_ratio = non_printable_count / len(text)

    # Determine if the text is mostly binary
    return binary_ratio > threshold


def remove_initial_all_caps(text):
    """Removes the initial all-caps section from the given text.

    Args:
      text: The input text.

    Returns:
      The text without the initial all-caps section.
    """

    # Your string to search in

    # Regular expression pattern
    pattern = r"[A-Z0-9 :,-]+"

    # Finding all matches
    matches = re.findall(pattern, text)

    if matches:
        # Getting the longest match
        longest_match = max(matches, key=len)

        # Finding the position of the longest match
        match_start = text.find(longest_match)

        # Deleting everything before and including the longest match
        result = text[match_start + len(longest_match) :]
    else:
        result = text
    return result


# Example usage

if __name__ == "__main__":
    path = "chat_server/chat/file_resources/cullman-municode/CODE_OF_ORDINANCES_CITY_OF_CULLMAN__ALABAMA.docx"
    chunk_municode_doc_into_paragraphs(path)
    # asdf = "DIVISION 1. GENERALLY"
    # doctracker = DocumentTracker()
    # doctracker.parse_input(asdf)
    # print(doctracker.current_location())
    # chuns = chunk_docx_into_paragraphs(path, 3000)
    # print(chuns)
    # print(len(chuns))
    # pdf_path = "chat_server/chat/file_resources/murray-muni-resources/https:$$$$$$codelibrary.amlegal.com$$$codes$$$murrayut$$$latest$$$murray_ut$$$0-0-0-14#codecontent.pdf"
    # pdf_path = "asdfasdf/test.pdf"
    # pdf_path = "pathpath/https:$$$$$$www.murray.utah.gov$$$ArchiveCenter$$$ViewFile$$$Item$$$87.pdf"
    # pdf_path = "pathpath/https:$$$$$$codelibrary.amlegal.com$$$codes$$$murrayut$$$latest$$$murray_ut$$$0-0-0-4630.pdf"
    # pdf_path = "pathpathpath/https:$$$$$$codelibrary.amlegal.com$$$codes$$$murrayut$$$latest$$$murray_ut$$$0-0-0-4630.pdf"
    # # Extract text from PDF
    # # extracted_text = chunk_docx_into_paragraphs(pdf_path)
    # validity = check_validity_of_file(pdf_path)
    # print(validity)
# main()
# chunk_docx_into_paragraphs("test.docx")


# chunk_pdf_into_paragraphs(pdf_path, "end_sentence_newline")
# xl_path = "chat_server/chat/file_resources/gunnison-resources/https:$$$$$$www.gunnisoncounty.org$$$DocumentCenter$$$View$$$14027$$$Public-Data---VALUES-2824.xlsx"
# chunks = chunk_xlsx_into_paragraphs(xl_path)
# for i in range(10):
#     print(chunks[i])
