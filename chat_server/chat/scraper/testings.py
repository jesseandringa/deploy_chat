import os
import re
import time
from base64 import b64decode

import file_reader
import html_helper
import requests
import website_scraper
from docx import Document
from selenium import webdriver


def main():
    # url = "https://www.murray.utah.gov/ArchiveCenter/ViewFile/Item/87"
    url = "https://codelibrary.amlegal.com/codes/murrayut/latest/murray_ut/0-0-0-4630"
    driver = webdriver.Chrome()
    # save_page_as_pdf(driver, url)
    get_text(driver, url)
    # os.makedirs("asdfasdf", exist_ok=True)
    # download_pdf(url, "asdfasdf/test.pdf")


def download_pdf(url, file_path):
    response = requests.get(url)
    with open(file_path, "wb") as file:
        file.write(response.content)


def make_folder_from_current_path(folder_name):
    os.makedirs(folder_name, exist_ok=True)


def save_page_as_pdf(driver, url):
    filename = website_scraper.convert_url_to_file_name(url)
    print("hello", url, filename)
    save_path = "pathpath" + "/" + filename + ".pdf"
    # make folder if not exists: pathpath
    make_folder_from_current_path("pathpath")
    try:
        html_helper.write_webpage_to_pdf(None, url, save_path)
        # print("saved pdf")
    except Exception:
        # print("couldnnt save pdf", e)
        pass
    # return save_path
    driver.get(url)
    pdf_data = driver.execute_cdp_cmd(
        "Page.printToPDF",
        {
            "path": "pathpathpath"
            + "/"
            + website_scraper.convert_url_to_file_name(url)
            + ".pdf",
            "format": "A4",
        },
    )
    make_folder_from_current_path("pathpathpath")
    with open(
        "pathpathpath" + "/" + website_scraper.convert_url_to_file_name(url) + ".pdf",
        "wb",
    ) as f:
        f.write(b64decode(pdf_data["data"]))
    # print("done:::  ", pdf_data["data"])
    chunks = file_reader.chunk_pdf_into_paragraphs(save_path)
    print("chunks", chunks)


def get_text(driver, url):
    from bs4 import BeautifulSoup

    # Set up the WebDriver (e.g., Chrome)
    # driver = webdriver.Chrome()
    driver.implicitly_wait(10)
    # Load the page
    driver.get(url)
    time.sleep(3)
    # Get the page source
    html = driver.page_source

    # Parse the HTML with BeautifulSoup
    soup = BeautifulSoup(html, "html.parser")

    # Extract all the text
    page_text = soup.get_text()
    print(page_text)
    text = clean_text(page_text)
    # print(text)
    create_docx(str(text), "test.docx")
    read_docs("test.docx")

    # Close the browser
    driver.quit()


def clean_text(text, max_length=100):
    # Define a regex pattern to find long continuous strings without spaces
    pattern = r"\S{%d,}" % max_length
    # Replace matched patterns with empty string
    cleaned_text = re.sub(pattern, "", text)
    return cleaned_text


def create_docx(text, filename):
    doc = Document()
    p = doc.add_paragraph()
    p.add_run(text)
    doc.save(filename)


def read_docs(filename):
    # Open the .docx file
    doc = Document(filename)
    extracted_text = ""
    for para in doc.paragraphs:
        print("\n new line: \n", para.text)
        extracted_text += para.text + "\n"
    return extracted_text


if __name__ == "__main__":
    # main()
    read_docs("test.docx")
