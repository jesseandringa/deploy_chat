import os
import re
import time
from base64 import b64decode

import html_helper
import requests
from bs4 import BeautifulSoup
from docx import Document
from urllib3.exceptions import MaxRetryError


def is_valid_xlsx(file_path):
    import openpyxl

    """Check if the file at `file_path` is a valid XLSX file.

    Args:
        file_path (str): The path to the XLSX file to check.

    Returns:
        bool: True if the file is a valid XLSX file, False otherwise.
    """
    try:
        workbook = openpyxl.load_workbook(file_path)
        return True
    except:
        return False


def convert_url_to_file_name(url):
    path = url.replace("/", "$$$")
    url = path
    return url


def download_as_xlsl(driver, url, file_path):
    filename = convert_url_to_file_name(url) + ".xlsx"
    # Construct the full save path
    save_path = os.path.join(file_path, filename)

    # Send a GET request to download the file
    response = requests.get(url, stream=True)
    # Check for successful response
    if response.status_code == 200:
        # Open the file in binary write mode
        with open(save_path, "wb") as file:
            for chunk in response.iter_content(1024):
                # Write the downloaded data in chunks
                file.write(chunk)
            if not is_valid_xlsx(save_path):
                os.remove(save_path)
                # print(f"Error downloading file: {response.status_code}")
                # print(save_path)

    return ".xlsx"


def download_page_with_requests(driver, url, file_path):
    response = requests.get(url)
    just_file_name = convert_url_to_file_name(url)
    filename = file_path + "/" + just_file_name + ".pdf"
    with open(filename, "wb") as file:
        file.write(response.content)
    return ".pdf"


def download_page_using_wkhtmltopdf(driver, url, file_path):
    filename = convert_url_to_file_name(url)
    try:
        html_helper.write_webpage_to_pdf(driver, url, file_path)
        # print("saved pdf")
    except Exception:
        # print("couldnnt save pdf", e)
        pass
    return ".pdf"


def download_page_using_chrome_print(driver, url, file_path):
    filename = convert_url_to_file_name(url)

    try:
        driver.get(url)
    except MaxRetryError:
        return ".pdf"

    pdf_data = driver.execute_cdp_cmd(
        "Page.printToPDF",
        {
            "path": file_path + "/" + filename + ".pdf",
            "format": "A4",
        },
    )
    with open(
        file_path + "/" + filename + ".pdf",
        "wb",
    ) as f:
        f.write(b64decode(pdf_data["data"]))

    return ".pdf"


def download_page_using_pagesource_to_docx(
    driver,
    url,
    filepath,
):
    driver.implicitly_wait(10)
    # Load the page
    driver.get(url)
    time.sleep(3)
    # Get the page source
    html = driver.page_source

    # Parse the HTML with BeautifulSoup
    soup = BeautifulSoup(html, "html.parser")

    # Extract all the text
    page_text = soup.get_text()
    # print(page_text)
    text = clean_text(page_text)
    filename = filepath + "/" + convert_url_to_file_name(url) + ".docx"
    # print(text)
    create_docx(str(text), filename)

    # Close the browser
    return ".docx"


def clean_text(text, max_length=100):
    # Define a regex pattern to find long continuous strings without spaces
    pattern = r"\S{%d,}" % max_length
    # Replace matched patterns with empty string
    text = re.sub(pattern, "", text)
    text = text.replace("\n\n", " ")
    start_marker = "Afrikaans"
    end_marker = "Zulu"
    start_position = text.find(start_marker)
    end_position = text.find(end_marker, start_position)

    if start_position != -1 and end_position != -1:
        # Include the length of the end_marker to remove it completely
        text = (
            text[:start_position].strip()
            + text[end_position + len(end_marker) :].strip()
        )

    return text


def create_docx(text, filename):
    doc = Document()
    p = doc.add_paragraph()
    p.add_run(text)
    doc.save(filename)


# import file_reader
# from selenium import webdriver
# from selenium.webdriver.chrome.options import Options as ChromeOptions


# def download_file(url, directory, is_html_page=False):
#     file_path = directory + "/" + convert_url_to_file_name(url)
#     print("file_path: ", file_path)
#     chrome_options = ChromeOptions()
#     chrome_options.add_argument("--headless")
#     driver = webdriver.Chrome(chrome_options)

#     os.makedirs(directory, exist_ok=True)
#     # filename = convert_url_to_file_name(url) + extension
#     validity = []
#     if is_html_page:
#         funcs = [
#             # downloader.download_page_with_requests,
#             download_page_using_chrome_print,
#             # downloader.download_page_using_wkhtmltopdf,
#             download_page_using_pagesource_to_docx,
#         ]
#     else:
#         funcs = [
#             download_page_with_requests,
#             download_page_using_chrome_print,
#             # downloader.download_page_using_wkhtmltopdf,
#             download_page_using_pagesource_to_docx,
#             download_as_xlsl,
#         ]
#     pdfs = file_reader.get_all_pdfs(".")
#     num_pdfs_before = len(pdfs)
#     for i in range(len(funcs)):
#         # print("trying: ", str(funcs[i]))
#         try:
#             extension = funcs[i](driver, url, directory)
#         except Exception:
#             extension = ".xlsx"
#         temp_file = file_path + extension
#         try:
#             validity.append(file_reader.check_validity_of_file(temp_file))
#         except Exception:
#             validity.append(0)
#     # print("validity: ", validity)
#     max_value = max(validity)
#     delete_file(file_path + ".pdf")
#     delete_file(file_path + ".docx")
#     delete_file(file_path + ".xlsx")

#     pdfs = file_reader.get_all_pdfs(".")
#     num_pdfs_after = len(pdfs)
#     if num_pdfs_after > num_pdfs_before:
#         for i in range(num_pdfs_after):
#             if pdfs[i].lower() in file_path.lower():
#                 dir_name = os.path.dirname(pdfs[i])
#                 new_path = os.path.join(dir_name, file_path + ".pdf")
#                 os.rename(pdfs[i], new_path)
#                 return
#     index = 0
#     if max_value > 0:
#         index = validity.index(max_value)
#         funcs[index](driver, url, directory)


# def delete_file(filepath):
#     try:
#         os.remove(filepath)
#     except FileNotFoundError:
#         pass
#     except PermissionError:
#         pass


# url = "https://cdn.kingcounty.gov/-/media/king-county/depts/council/maps/2021_districting_plan_20211208-pdf.pdf?rev=4e60b87c17cb4231bdea766e23637963&hash=6B895499CD11D69D73A941A15AA10F5F"
# download_file(url, "chat_server/chat/file_resources/king-wa-resources", False)
# driver = webdriver.Chrome()
# url = "https://codelibrary.amlegal.com/codes/murrayut/latest/murray_ut/0-0-0-2382"
# file_path = "chat_server"
# download_page_using_pagesource_to_docx(driver, url, file_path)
